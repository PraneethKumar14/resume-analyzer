from docx import Document
import io

def extract_text_from_docx(file_stream):
    """Extract text from DOCX file with better error handling"""
    try:
        file_stream.seek(0)
        doc = Document(file_stream)
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract from tables (important for many resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        cleaned_text = clean_docx_text(text)
        print(f"✅ Successfully extracted {len(cleaned_text)} characters from DOCX")
        return cleaned_text
        
    except Exception as e:
        print(f"❌ Error extracting DOCX text: {e}")
        return "ERROR: Could not extract text from DOCX"

def clean_docx_text(text):
    """Clean and normalize DOCX text"""
    if not text:
        return ""
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped and len(stripped) > 1:
            cleaned_lines.append(stripped)
    return '\n'.join(cleaned_lines)
                    